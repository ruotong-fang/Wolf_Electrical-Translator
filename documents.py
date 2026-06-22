from dataclasses import dataclass
from pathlib import Path
import codecs
from typing import Callable, Dict, List, Optional


class DocumentError(RuntimeError):
    pass


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    kind: str
    warning: str = ""
    encoding: str = ""


@dataclass(frozen=True)
class TextCandidate:
    encoding: str
    text: str
    confidence: float


BOM_ENCODINGS = (
    (codecs.BOM_UTF8, "utf-8-sig"),
    (codecs.BOM_UTF32_LE, "utf-32"),
    (codecs.BOM_UTF32_BE, "utf-32"),
    (codecs.BOM_UTF16_LE, "utf-16"),
    (codecs.BOM_UTF16_BE, "utf-16"),
)


def _text_quality(text: str) -> float:
    if not text:
        return 0.0
    controls = sum(ord(char) < 32 and char not in "\r\n\t" for char in text)
    replacements = text.count("\ufffd")
    nulls = text.count("\x00")
    suspicious = controls + replacements * 8 + nulls * 8
    return max(0.0, 1.0 - suspicious / max(len(text), 1))


def text_candidates(path: Path) -> List[TextCandidate]:
    raw = path.read_bytes()
    if not raw:
        return [TextCandidate("utf-8", "", 1.0)]
    for bom, encoding in BOM_ENCODINGS:
        if raw.startswith(bom):
            return [TextCandidate(encoding, raw.decode(encoding), 1.0)]

    candidates: Dict[str, TextCandidate] = {}
    encodings = ("utf-8", "gb18030", "big5", "utf-16-le", "utf-16-be", "windows-1252")
    for encoding in encodings:
        try:
            text = raw.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
        quality = _text_quality(text)
        if encoding.startswith("utf-16") and b"\x00" not in raw:
            quality *= 0.45
        if encoding == "windows-1252" and any(byte >= 0x80 for byte in raw):
            quality *= 0.7
        candidates[encoding] = TextCandidate(encoding, text, quality)

    try:
        from charset_normalizer import from_bytes

        match = from_bytes(raw).best()
        if match and match.encoding:
            encoding = match.encoding.lower()
            text = str(match)
            confidence = max(0.0, 1.0 - float(match.chaos)) * _text_quality(text)
            current = candidates.get(encoding)
            if current is None or confidence > current.confidence:
                candidates[encoding] = TextCandidate(encoding, text, confidence)
    except (ImportError, UnicodeError, ValueError):
        pass

    if not candidates:
        raise DocumentError("无法识别文本文件编码")
    return sorted(candidates.values(), key=lambda item: item.confidence, reverse=True)


def _read_text(path: Path, encoding: Optional[str] = None) -> ExtractedDocument:
    if encoding:
        try:
            return ExtractedDocument(path.read_text(encoding=encoding), "文本", encoding=encoding)
        except (UnicodeDecodeError, LookupError) as exc:
            raise DocumentError(f"无法使用 {encoding} 读取这个文件") from exc
    candidates = text_candidates(path)
    best = candidates[0]
    warning = ""
    if best.confidence < 0.75 or (len(candidates) > 1 and best.confidence - candidates[1].confidence < 0.08):
        warning = f"文本编码可能是 {best.encoding}，请检查预览内容"
    return ExtractedDocument(best.text, "文本", warning, best.encoding)


def _read_pdf(path: Path) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentError("当前应用未包含 PDF 组件，请安装完整版应用") from exc

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentError(f"PDF 读取失败：{exc}") from exc
    text = "\n\n".join(page.strip() for page in pages if page.strip())
    if not text:
        raise DocumentError("这个 PDF 没有可提取文字，可能是扫描件，需要 OCR 版本处理")
    return ExtractedDocument(text, "PDF", "已提取 PDF 文字；复杂表格和版面可能需要人工核对")


def _read_docx(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentError("当前应用未包含 Word 组件，请安装完整版应用") from exc

    try:
        document = Document(str(path))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text.strip() for cell in row.cells))
    except Exception as exc:
        raise DocumentError(f"Word 读取失败：{exc}") from exc
    return ExtractedDocument("\n".join(blocks), "Word", "已提取正文和表格文字")


def _read_xlsx(path: Path) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentError("当前应用未包含 Excel 组件，请安装完整版应用") from exc

    try:
        workbook = load_workbook(str(path), read_only=True, data_only=False)
        blocks = []
        for sheet in workbook.worksheets:
            blocks.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    blocks.append("\t".join(values))
        workbook.close()
    except Exception as exc:
        raise DocumentError(f"Excel 读取失败：{exc}") from exc
    return ExtractedDocument("\n".join(blocks), "Excel", "已提取单元格内容；公式和原格式不会写入译文")


READERS: Dict[str, Callable[[Path], ExtractedDocument]] = {
    ".txt": _read_text,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
}


def extract_document(path: str, encoding: Optional[str] = None) -> ExtractedDocument:
    file_path = Path(path)
    reader = READERS.get(file_path.suffix.lower())
    if reader is None:
        raise DocumentError("暂不支持这个文件类型，请选择 TXT、PDF、DOCX 或 XLSX")
    if not file_path.is_file():
        raise DocumentError("文件不存在")
    if file_path.suffix.lower() == ".txt":
        return _read_text(file_path, encoding)
    return reader(file_path)
